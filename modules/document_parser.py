from pathlib import Path

from docx import Document
from pypdf import PdfReader

SCAN_WARNING = "Текст из PDF не извлечен. Вероятно, документ является сканированным. Для анализа требуется текстовый PDF, DOCX или TXT."


def _extract_pdf_with_pypdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def _extract_pdf_with_pdfplumber(file_path: Path) -> str:
    try:
        import pdfplumber
    except Exception:
        return ""
    with pdfplumber.open(str(file_path)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()


def extract_text(path: str | Path) -> tuple[str, str | None]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".txt":
        return file_path.read_text(encoding="utf-8", errors="ignore"), None
    if suffix == ".docx":
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_cells = []
        for table in doc.tables:
            for row in table.rows:
                table_cells.extend(cell.text for cell in row.cells if cell.text.strip())
        return "\n".join(paragraphs + table_cells), None
    if suffix == ".pdf":
        text = _extract_pdf_with_pypdf(file_path) or _extract_pdf_with_pdfplumber(file_path)
        if not text:
            return "", SCAN_WARNING
        return text, None
    return "", "Поддерживаются только DOCX, PDF и TXT."
