from pathlib import Path
from docx import Document
from .database import EXPORT_DIR

SECTIONS = [("Краткий вывод", "summary"), ("Что добавлено", "added_items"), ("Что исключено", "removed_items"), ("Что изменилось по смыслу", "semantic_changes"), ("Новые обязанности", "new_obligations"), ("Новые права", "new_rights"), ("Запреты и ограничения", "prohibitions"), ("Финансовые последствия", "financial"), ("Операционное влияние", "operations"), ("Полномочия государственных органов", "authority"), ("Риск расширительного толкования", "uncertainty"), ("Итоговый юридический обзор влияния на деятельность Общества", "final_review")]

def _add_table(doc, rows: list[dict], columns: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns)); table.style = "Table Grid"
    for i, col in enumerate(columns): table.rows[0].cells[i].text = col
    for row in rows:
        cells = table.add_row().cells
        for i, col in enumerate(columns): cells[i].text = str(row.get(col, ""))

def export_analysis_to_word(result: dict, user_login: str) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe = (result.get("document_title") or "analysis").replace("/", "_").replace("\\", "_")[:60]
    path = EXPORT_DIR / f"legal_review_{safe}.docx"
    doc = Document(); doc.add_heading("Legal Redline Analyzer", 0)
    doc.add_paragraph(f"Название анализа: {result.get('document_title', '')}")
    doc.add_paragraph(f"Дата: {result.get('created_at', '')}")
    doc.add_paragraph(f"Пользователь: {user_login}")
    for title, key in SECTIONS:
        doc.add_heading(title, level=1)
        value = result.get(key, "")
        if isinstance(value, list):
            for item in value: doc.add_paragraph(str(item), style="List Bullet")
        else:
            for part in str(value).split("\n"): doc.add_paragraph(part)
    doc.add_heading("Сравнение редакций", level=1)
    _add_table(doc, result.get("changes", []), ["№", "Элемент", "Действующая редакция", "Предлагаемая редакция", "Тип изменения"])
    doc.add_heading("Таблица рисков", level=1)
    _add_table(doc, result.get("risks", []), ["№", "Риск", "Описание", "Уровень риска", "Возможное влияние на Общество", "Комментарий"])
    doc.save(path)
    return path
